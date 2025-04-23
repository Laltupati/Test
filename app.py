import os
import re
import requests
import json
from flask import Flask, render_template, request, jsonify, send_file, session
from azure.storage.blob import generate_blob_sas, BlobSasPermissions
from flask_session import Session
from PyPDF2 import PdfReader
import docx,urllib.parse
from collections import Counter
from docx import Document
from docx.shared import Inches  
from flask.helpers import send_from_directory
from azure.storage.blob import ContainerClient
from azure.storage.blob import BlobServiceClient, BlobClient
from io import BytesIO
import tempfile
import docx2txt
import io
from uuid import uuid4
from datetime import datetime, timedelta
from flask.sessions import SessionInterface, SessionMixin
from werkzeug.datastructures import CallbackDict
# Suppress InsecureRequestWarning
import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

app = Flask(__name__)
app.secret_key = "super_secret_key"
# AZURE_STORAGE_CONNECTION_STRING = "DefaultEndpointsProtocol=https;AccountName=;AccountKey=;EndpointSuffix=core.windows.net"
UPLOAD_CONTAINER_NAME = "uploaded-files"
ACCOUNT_NAME = ""
ACCOUNT_KEY = ""
AZURE_STORAGE_CONNECTION_STRING = f"DefaultEndpointsProtocol=https;AccountName={ACCOUNT_NAME};AccountKey={ACCOUNT_KEY};EndpointSuffix=core.windows.net"
# SESSION_CONTAINER_NAME = "session-data"
class AzureBlobStorageSession(CallbackDict, SessionMixin):
    def __init__(self, container_client, session_id):
        self.container_client = container_client
        self.session_id = session_id
        data = {}
        try:
            blob_client = container_client.get_blob_client(session_id)
            blob_data = blob_client.download_blob().content_as_text()
            data = json.loads(blob_data)
        except Exception:
            pass
        super(AzureBlobStorageSession, self).__init__(data)

class AzureBlobStorageSessionInterface(SessionInterface):
    def __init__(self, container_name):
        self.container_name = container_name
    def _get_container_client(self):
        return ContainerClient.from_connection_string(conn_str=AZURE_STORAGE_CONNECTION_STRING, container_name=self.container_name)
    def get_expiration_time(self, app, session):
        return datetime.utcnow() + timedelta(seconds=app.config['PERMANENT_SESSION_LIFETIME'])
    def get_session_id(self):
        return f"session_{uuid4()}"
    def open_session(self, app, request):
        session_id = request.cookies.get(app.session_cookie_name)
        container_client = self._get_container_client()
        return AzureBlobStorageSession(container_client, session_id)
    def save_session(self, app, session, response):
        expiration_time = self.get_expiration_time(app, session)
        session_id = session.session_id if session.session_id else self.get_session_id()
        container_client = self._get_container_client()
        blob_client = container_client.get_blob_client(session_id)
        blob_data = json.dumps(dict(session))
        blob_client.upload_blob(blob_data, blob_type='BlockBlob', overwrite=True)
        response.set_cookie(app.session_cookie_name, session_id, expires=expiration_time, httponly=True, domain=self.get_cookie_domain(app))
# app.config["PERMANENT_SESSION_LIFETIME"] = 86400  # Duration in seconds (one day)
# app.session_interface = AzureBlobStorageSessionInterface(SESSION_CONTAINER_NAME)
# Session(app)
ALLOWED_EXTENSIONS = {"pdf", "docx"}
QUALITIES = ['communication', 'programming', 'leadership', 'teamwork', 'analytical', 'creative']
def allowed_file(filename):
    return (
        "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
    )
def get_container_client(container_name):
    return ContainerClient.from_connection_string(conn_str=AZURE_STORAGE_CONNECTION_STRING, container_name=container_name)
def upload_file_to_blob(file, container_name):
    container_client = get_container_client(container_name)
    blob_client = container_client.get_blob_client(file.filename)
    with BytesIO() as file_stream:
        file.save(file_stream)
        file_stream.seek(0)
        blob_client.upload_blob(file_stream, overwrite=True)
    return blob_client.url
def download_blob_to_text(blob_url, is_binary=False):
    response = requests.get(blob_url)
    if is_binary:
        return response.content
    else:
        return response.text
def extract_resume_skill_scores(resume_lines):
    total_skill_score = 0
    for line in resume_lines:
        match = re.match(r"\s*\|\s*[^|]+\s*\|\s*\d+\s*\|\s*(\d+)\s*\|", line)
        if match:
            skill_score = int(match.group(1))
            total_skill_score += skill_score
    return total_skill_score
def get_blob_sas(account_name, account_key, container_name, blob_name):
    sas_token = generate_blob_sas(
        account_name=account_name,
        account_key=account_key,
        container_name=container_name,
        blob_name=blob_name,
        permission=BlobSasPermissions(read=True),  # set permissions according to your requirement
        expiry=datetime.utcnow() + timedelta(hours=1)  # SAS token will expire 1 hour from now
    )
    return sas_token
def read_document_content(blob_url):
    resume_file_name = os.path.basename(blob_url)
    sas_token = get_blob_sas(ACCOUNT_NAME, ACCOUNT_KEY, UPLOAD_CONTAINER_NAME, urllib.parse.unquote(resume_file_name))
    blob_url = f"https://{ACCOUNT_NAME}.blob.core.windows.net/{UPLOAD_CONTAINER_NAME}/{resume_file_name}?{sas_token}"
    # blob_url = f"{blob_url}?{sas_token}"
    response = requests.get(blob_url)
    file_data = response.content
    # print(f"file_data: {file_data}")
    if resume_file_name.endswith(".pdf"):
        with BytesIO(file_data) as f:
            try:
                pdf_reader = PdfReader(f)
                extracted_text = ""
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text()
            except Exception as e:
                # print(f"Error while reading PDF file: {e}")
                extracted_text = ""
        return extracted_text
    elif resume_file_name.endswith(".docx"):
        try:
            document = docx.Document(io.BytesIO(file_data))
            extracted_text = ""
            for paragraph in document.paragraphs:
                extracted_text += paragraph.text + "\n"
        except Exception as e:
            # print(f"Error while reading .docx file: {e}")
            extracted_text = ""
        return extracted_text
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    resume_files = request.files.getlist("resumes")
    jd_file = request.files["jd"]
    jd_filepath = upload_file_to_blob(jd_file, UPLOAD_CONTAINER_NAME)
    uploaded_resumes = []
    for file in resume_files:
        if file and allowed_file(file.filename):
            file_url = upload_file_to_blob(file, UPLOAD_CONTAINER_NAME)
            uploaded_resumes.append({"name": file.filename, "path": file_url})
    # Store uploaded resumes and JD in session
    session["resumes"] = uploaded_resumes
    session["jd"] = jd_filepath
    return jsonify({"resumes": uploaded_resumes, "jd": jd_filepath})

@app.route("/api/send_prompt", methods=["POST"])
def send_prompt():
    prompt = request.form.get("prompt")
    conversation_history = request.form.get("conversation_history")
    
    api_url = ""
    api_key = "052bc9aaeb0b4ae6bccc42ef4d05e1b5"
    headers = {"api-key": api_key}
    query_data = {
        "messages": [
                {"role": "system", "content": f"{conversation_history}\n{prompt}"}
            ],
            "temperature": 0.3,
            "top_p": 1,
            "frequency_penalty": 0,
            "presence_penalty": 0,
            "max_tokens": 4000,
            "stop": None,
        }
    
    response = requests.post(api_url, headers=headers, json=query_data, verify=False)
    response_data = response.json()
    generated_text = response_data["choices"][0]["message"]["content"].strip()
    # print(generated_text)
    return jsonify({"response": generated_text})
def send_gpt3_request(prompt):
    api_url = ""
    api_key = ""
    headers = {"api-key": api_key}
    query_data = {
        "messages": [
                {"role": "system", "content": prompt}
            ],
            "temperature": 0.3,
            "top_p": 1,
            "frequency_penalty": 0,
            "presence_penalty": 0,
            "max_tokens": 4000,
            "stop": None,
        }
    
    try:
        response = requests.post(api_url, headers=headers, json=query_data, verify=False)
        response_data = response.json()
        if "choices" not in response_data:
            print("Error response from the GPT-3 API:", response_data)
            return {"choices": [{"text": ""}]}
        return response_data
    
    except Exception as e:
        print(f"Error while sending request to GPT-3 API: {e}")
        return {"choices": [{"text": ""}]}
def extract_skills_score(generated_text):
    match = re.match(
        r"\s*Total Skillset Score:\s*(\d+)",
        generated_text,
    )
    if match:
        return int(match.group(1))
    return 0

@app.route("/api/display_skills_ranking_table", methods=["POST"])
def display_skills_ranking_table():
    resumes = session.get("resumes")
    jd = session.get("jd")
    jd_file_name = os.path.basename(jd)
    resume_file_names = [os.path.basename(resume["path"]) for resume in resumes]
    prompt = f"Please provide the total skillset scores summarizing the compatibility of each of these resumes with the job requirements for {jd_file_name}. The resumes are: {', '.join(resume_file_names)}. Provide the scores in the following format: \"Resume: {{resume_file_name}} - Total Skillset Score: {{score}}\"."
    conversation_history_with_prompt = f"\n{prompt}"
    response_data = send_gpt3_request(conversation_history_with_prompt)
    generated_text = response_data["choices"][0]["message"]["content"].strip()
    lines = generated_text.split("\n")
    data = [(m.group(1), int(m.group(2))) for line in lines if (m := re.match(r"\s*Resume:\s*(.+)\s*-\s*Total Skillset Score:\s*(\d+)", line)) is not None]
    html_table = convert_to_html_table(data)
    session["skill_ranking_table"] = html_table  # Add this line to save the HTML table in the session
    return jsonify({"table": html_table})
@app.route("/api/display_pros_cons_table", methods=["POST"])
def display_pros_cons_table():
    resumes = session.get("resumes")
    jd = session.get("jd")
    results = []
    for index, resume in enumerate(resumes, start=1):
        resume_file_name = os.path.basename(resume["path"])
        prompt = f"Display the pros and cons of candidate {index}: {resume_file_name}. Use the format: \"Pros: {{pros}}; Cons: {{cons}}\"."
        conversation_history_with_prompt = f"{read_document_content(resume['path'])}\n{read_document_content(jd)}\n{prompt}"
        # print (f">>>>>conversation history {conversation_history_with_prompt}")
        response_data = send_gpt3_request(conversation_history_with_prompt)
        generated_text = response_data["choices"][0]["message"]["content"].strip()
        formatted_text = f"Candidate {index}: {urllib.parse.unquote(resume_file_name)} - {generated_text}\n"
        results.append(formatted_text)
    # Format the pros and cons text and store it in the session
    formatted_text_list = []
    for result in results:
        formatted_text = result.replace("Pros:", "<br>Pros:").replace("Cons:", "<br>Cons:")
        formatted_text += "<hr>"
        formatted_text_list.append(formatted_text)
    session["pros_cons_html"] = "".join(formatted_text_list)
    return jsonify({"results": results})

def extract_pros_and_cons(generated_text):
    match = re.match(r"\s*Pros:\s*(.+)\s*;\s*Cons:\s*(.+)", generated_text)
    if match:
        return match.groups()
    return "", ""
@app.route("/api/send_custom_prompt", methods=["POST"])
def send_custom_prompt():
    prompt = request.form.get("prompt")
    resumes_text = "\n".join([read_document_content(resume["path"]) for resume in session.get("resumes")])
    conversation_history = request.form.get("conversation_history", "")
    skill_ranking_table_html = session.get("skill_ranking_table", "")
    pros_cons_html = session.get("pros_cons_html", "")
    prompt_with_content = (f"{resumes_text}\n"
                           f"{read_document_content(session['jd'])}\n"
                           f"{skill_ranking_table_html}\n"
                           f"{pros_cons_html}\n"
                           f"{conversation_history}\n"
                           f"{prompt}\n"
                           "Stay in character and don't provide generic responses. Provide quantification whenever possible.")
    
    response_data = send_gpt3_request(prompt_with_content)
    generated_text = response_data["choices"][0]["message"]["content"].strip()
    return jsonify({"response": generated_text})
def convert_to_html_table(data):
    html_table = '<table border="1" cellpadding="5" cellspacing="0">'
    html_table += "<tr><th colspan='2'>Skills Ranking Table</th></tr>"
    html_table += "<tr><th>Resume</th><th>Skills Score</th></tr>"
    for row in data:
        html_table += f"<tr><td>{row[0]}</td><td>{row[1]}</td></tr>"
    html_table += "</table>"
    explanation = (
        "<p>The skill scores are calculated based on an AI-powered analysis of the content "
        "of the uploaded resumes and the job description. GPT-3.5 Turbo, a powerful natural "
        "language processing model, is used to estimate each candidate's compatibility "
        "with the job requirements by factoring in various skills and qualifications "
        "mentioned in their resumes. Please note that these scores are meant to "
        "provide an initial assessment and should not replace human judgment in the "
        "final decision-making process.</p>"
    )
    html_table += explanation
    return html_table
@app.route("/api/export_output_in_ms_word", methods=["POST"])
def export_output_in_ms_word():
    content = request.form.get("content")
    document = Document()
    p = document.add_paragraph(content)
    # Save the document
    output_file = "output.docx"
    document.save(output_file)
    # Send the file to the user
    return send_file(output_file, as_attachment=True, download_name=output_file)

# Run the application
if __name__ == "__main__":
    app.run(debug=True)
